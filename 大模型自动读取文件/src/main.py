# -*- coding: utf-8 -*-
import os
import sys
import glob
import subprocess
import docx
from doc_handler import DocumentHandler
from llm_client import LLMClient
from rag_engine import RAGEngine

def fill_student_info(doc):
    """
    自动填入默认的学生及课程基础信息（学号姓名等），省去手动修改的麻烦。
    """
    info_mapping = {
        "课程编号": "03237041",
        "课程名称": "Python语言程序设计",
        "学院": "计算机与人工智能",
        "专业": "人工智能",
        "学号": "114514",
        "姓名": "龙傲天"
    }
    
    # 遍历 Word 中的所有表格
    for table in doc.tables:
        for row in table.rows:
            # 1. 检查“键-值”分开在相邻两个单元格的情况（最常见，如 左边单元格是“姓名：”，右边是“陈佳翔”）
            for c_idx in range(len(row.cells) - 1):
                cell_left = row.cells[c_idx]
                cell_right = row.cells[c_idx + 1]
                
                # 移除空格干扰进行比较
                left_text = cell_left.text.replace(" ", "").replace("　", "")
                for key, val in info_mapping.items():
                    if left_text == key or left_text == f"{key}:" or left_text == f"{key}：":
                        cell_right.text = val
                        
            # 2. 检查“键-值”合并在同一个单元格里的情况（如：“学号：2515308151”）
            for cell in row.cells:
                cell_text = cell.text
                for key, val in info_mapping.items():
                    if key in cell_text and val not in cell_text:
                        for sep in ("：", ":"):
                            if sep in cell_text:
                                parts = cell_text.split(sep)
                                if key in parts[0].replace(" ", ""):
                                    parts[1] = val
                                    cell.text = sep.join(parts)
                                    break

def find_report_cells(doc):
    """
    动态定位实验要求、实验结果和实验代码对应的目标单元格。
    避免硬编码下标（如 Table 0, Row 3 等），极大地提升对不同实验报告 Word 模板的兼容性。
    """
    req_cell = None
    results_cell = None
    code_cell = None
    
    # 遍历 Word 文档中的每一个表格
    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # 1. 寻找“实验要求”或“实验目的”标题单元格，其下一行的单元格即是具体的实验要求内容
                if "实验目的" in cell_text or "实验要求" in cell_text:
                    if "结果" not in cell_text and "代码" not in cell_text:
                        if r_idx + 1 < len(table.rows):
                            req_cell = table.cell(r_idx + 1, c_idx)
                            
                # 2. 寻找“实验结果（截图）”标题单元格，其下一行是回填实验图表与文字分析的目标位置
                if "实验结果" in cell_text or "结果（截图）" in cell_text:
                    if r_idx + 1 < len(table.rows):
                        results_cell = table.cell(r_idx + 1, c_idx)
                        
                # 3. 寻找“实验代码”标题单元格，其下一行是回填源码的目标位置
                if "实验代码" in cell_text or "三、实验代码" in cell_text:
                    if r_idx + 1 < len(table.rows):
                        code_cell = table.cell(r_idx + 1, c_idx)
                        
    return req_cell, results_cell, code_cell

def main():
    print("=== 开始 Python 实验报告自动生成与排版优化任务 ===")
    
    # 1. 初始化路径与目录
    raw_dir = "./data/raw"             # 存放您的旧报告（用于风格学习）
    templates_dir = "./data/templates" # 存放空白实验报告模板
    output_dir = "./output"             # 生成的最终报告输出目录
    temp_output_dir = "./data/output"   # 生成结果图片的临时目录
    
    os.makedirs(raw_dir, exist_ok=True)
    os.makedirs(templates_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(temp_output_dir, exist_ok=True)
    
    # 2. 自动检索空白实验报告 Word 模板文件，并让用户手动选择
    docx_files = glob.glob(os.path.join(templates_dir, "*.docx"))
    if not docx_files:
        print(f"❌ 错误：在 '{templates_dir}' 目录下未找到任何 .docx 格式的空白实验报告模板！")
        return
        
    print("\n-------------------------------------------")
    print("📁 检测到以下空白实验报告模板：")
    for idx, fpath in enumerate(docx_files):
        print(f"  [{idx + 1}] {os.path.basename(fpath)}")
    print("-------------------------------------------")
    
    # 💡 交互选择：让您指定要处理哪一份实验报告，防止自动化运行浪费 Token
    while True:
        try:
            user_input = input(f"请输入你想处理的报告序号 (1-{len(docx_files)}) [默认1]: ").strip()
            if not user_input:
                choice = 0
                break
            choice = int(user_input) - 1
            if 0 <= choice < len(docx_files):
                break
            else:
                print(f"⚠️ 请输入 1 到 {len(docx_files)} 之间的数字！")
        except ValueError:
            print("⚠️ 输入无效，请输入数字序号！")
            
    template_path = docx_files[choice]
    print(f"\n🚀 开始处理所选模板: {os.path.basename(template_path)}")
    
    # 3. 实例化处理器并解析所选的 Word 模板
    doc_handler = DocumentHandler(temp_dir=temp_output_dir)
    doc = docx.Document(template_path)
    
    # 💡 自动填入默认学生及课程基础信息（学号姓名等）
    print("📝 正在自动填入您的默认个人基本信息（学号、姓名、学院、专业等）...")
    fill_student_info(doc)
    
    # 动态定位关键单元格，防止硬编码下标越界或出错
    req_cell, results_cell, code_cell = find_report_cells(doc)
    
    if not req_cell:
        print("❌ 错误：未能动态定位到存放“实验目的与要求”的单元格，请检查模板！")
        return
    if not results_cell:
        print("❌ 错误：未能动态定位到回填“实验结果”的目标单元格，请检查模板！")
        return
    if not code_cell:
        print("❌ 错误：未能动态定位到回填“实验代码”的目标单元格，请检查模板！")
        return
        
    requirements = req_cell.text.strip()
    print(f"📌 成功定位并读取实验要求（共 {len(requirements)} 字）")
    
    # 4. 初始化 RAG 知识库，自动学习您的个人风格旧报告
    print("🔍 正在初始化 RAG 知识库引擎...")
    rag = RAGEngine()
    
    # 自动扫描个人风格文件目录并导入向量库
    style_files = []
    for ext in ("*.pdf", "*.docx", "*.txt", "*.md"):
        style_files.extend(glob.glob(os.path.join(raw_dir, ext)))
        
    if style_files:
        print(f"📄 检查到 {len(style_files)} 个您的文风样本。正在增量读取新加入的文件...")
        for style_file in style_files:
            # 内部会查重，仅在有新文件时产生 Embedding Token 消耗
            rag.ingest_file_via_mineru(style_file)
        print("💡 个人风格特征载入完成（重用本地缓存向量可达 0 Token 消耗）。")
    else:
        print("💡 未检测到个人风格样本。如果您希望生成的语气和您以往写的报告一样，可将几篇旧报告放入 `./data/raw` 目录。")
    
    # 5. 检索本地向量知识库，获取风格与背景参考材料
    print("🔍 正在检索 RAG 知识库以获取文风参考背景...")
    rag_context = rag.retrieve(requirements, top_k=2)
    if rag_context:
        print("📌 成功检索到相关风格参考，已注入大模型上下文。")
    else:
        print("⚠️ 未检索到高度相关的 RAG 参考资料，大模型将直接基于默认要求生成。")
    
    # 6. 调用大模型生成 Python 代码与分析文本（注入您的文风特征）
    llm = LLMClient()
    print("🤖 正在调用 Gemini 模型生成实验解析代码与文字分析...")
    solution = llm.generate_experiment_solution(requirements, rag_context=rag_context)
    
    code = solution.get("code", "")
    text_analysis = solution.get("text_analysis", "")
    
    if not code or "# 自动生成失败" in code:
        print("❌ 大模型生成方案失败，请检查 API Key 或网络连接！")
        return
    
    print("💡 实验代码与文字分析生成完毕。")
    
    # 7. 清理临时输出目录中的旧图片，防止混淆不同实验的结果图
    for old_img in glob.glob(os.path.join(temp_output_dir, "*.png")):
        try:
            os.remove(old_img)
        except Exception as e:
            print(f"⚠️ 无法删除旧图片 {old_img}: {e}")
            
    # 8. 为生成的代码包装安全运行“沙箱”，并在后台执行以产生结果截图
    current_workspace = os.path.abspath(os.getcwd())
    
    # 自动在代码顶部注入工作目录切换和 matplotlib 强制设置，防止相对路径找不到文件和绘图框弹窗
    defense_header = f"""# -*- coding: utf-8 -*-
import os
import sys
import matplotlib
matplotlib.use('Agg') # 强制使用非交互式后端，防止绘图时弹窗卡死程序
import matplotlib.pyplot as plt

# 强制切换当前运行路径为项目根目录，防止找不到 raw 里的数据文件
os.chdir(r"{current_workspace}")
sys.path.append(r"{current_workspace}")

# 设置全局中文字体与负号支持
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial']
plt.rcParams['axes.unicode_minus'] = False

# 确保图片输出目录存在
os.makedirs('./data/output', exist_ok=True)

# 以下为自动生成的实验解决方案代码：
"""
    wrapped_code = defense_header + "\n" + code
    
    temp_script_path = os.path.join(temp_output_dir, "temp_solution.py")
    with open(temp_script_path, "w", encoding="utf-8") as f:
        f.write(wrapped_code)
        
    print(f"💻 正在后台运行生成的 Python 代码以产生实验图表...")
    try:
        # 使用当前运行本程序的 Python 解释器在工作区根目录下执行代码
        result = subprocess.run(
            [sys.executable, temp_script_path],
            capture_output=True,
            text=True,
            cwd=current_workspace,
            encoding="utf-8",
            errors="ignore"
        )
        
        if result.returncode != 0:
            print("❌ 警告：生成的实验代码在后台运行时报错！错误日志如下：")
            print(result.stderr)
        else:
            print("🎉 实验代码运行成功，图表已生成。")
            if result.stdout.strip():
                print("--- 运行输出 ---")
                print(result.stdout)
                
    except Exception as e:
        print(f"❌ 运行生成的实验代码时发生异常: {e}")
        
    # 9. 自动搜集生成的所有实验图片路径
    image_paths = sorted(glob.glob(os.path.join(temp_output_dir, "*.png")))
    print(f"📊 收集到以下生成的实验结果图: {image_paths}")
    
    # 10. 将生成的代码与实验结果图片回填到 Word 文档对应位置
    print("📝 开始将代码与文字图片回填到实验报告中...")
    
    # 将实验分析和图片填在 results_cell (自动控制在 2 页内)
    doc_handler.write_results_to_cell(results_cell, text_analysis, image_paths)
    
    # 将代码填在 code_cell (自适应调整行距，控制在 3 页内)
    doc_handler.write_code_to_cell(code_cell, code)
        
    # 11. 保存最终优化排版后的实验报告
    output_filename = os.path.basename(template_path)
    final_output_path = os.path.join(output_dir, output_filename)
    
    doc.save(final_output_path)
    
    # 全局字体保修，将所有样式统一为五号宋体
    doc_handler.apply_styles_to_document(final_output_path)
    
    print(f"\n🎉 最终成果顺利生成！")
    print(f"👉 排版符合要求（全部五号宋体字，结果2页/代码3页内）")
    print(f"👉 优化后的实验报告已保存至: {final_output_path}")
    
    # 12. 清理临时运行脚本
    if os.path.exists(temp_script_path):
        try:
            os.remove(temp_script_path)
        except Exception:
            pass

if __name__ == "__main__":
    main()
