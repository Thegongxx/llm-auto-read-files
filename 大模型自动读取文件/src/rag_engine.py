# -*- coding: utf-8 -*-
import os
import json
import subprocess
import glob
import shutil
import docx
from google import genai
from google.genai import types
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

class RAGEngine:
    def __init__(self, db_path="./db/rag_store.json"):
        """
        初始化检索增强生成(RAG)向量知识库引擎
        """
        self.db_path = db_path
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        
        # 1. 统一从环境变量中获取 API Key
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("错误：未在环境变量中配置 GEMINI_API_KEY！请检查 .env 文件。")
            
        # 初始化 Gemini 客户端
        self.client = genai.Client(api_key=api_key)
        self.embed_model = "gemini-embedding-2"
        
        # 2. 从本地读取已经存好的向量知识库
        db_data = self._load_db()
        self.knowledge_base = db_data.get("chunks", [])
        self.ingested_files = db_data.get("ingested_files", [])

    def _load_db(self) -> dict:
        """
        内部辅助方法：从本地 JSON 加载知识库数据结构
        """
        if os.path.exists(self.db_path):
            try:
                with open(self.db_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    # 兼容之前只有列表的旧格式
                    if isinstance(data, list):
                        return {"chunks": data, "ingested_files": []}
                    return data
            except Exception as e:
                print(f"⚠️ 读取本地知识库 JSON 失败，将初始化为空库。原因: {e}")
        return {"chunks": [], "ingested_files": []}

    def save_db(self):
        """
        将当前内存中的向量和已导入文件名列表保存到本地 JSON 文件中
        """
        with open(self.db_path, "w", encoding="utf-8") as f:
            json.dump({
                "chunks": self.knowledge_base,
                "ingested_files": self.ingested_files
            }, f, ensure_ascii=False, indent=2)
        print(f"💾 知识库已成功保存到: {self.db_path}")

    def get_embedding(self, text: str) -> list:
        """
        调用官方 Gemini Embedding API 获取输入文本的向量表征（一串浮点数）
        """
        try:
            response = self.client.models.embed_content(
                model=self.embed_model,
                contents=text,
            )
            return response.embeddings[0].values
        except Exception as e:
            print(f"❌ 调用 Embedding API 获取向量失败: {e}")
            return []

    def add_document(self, text: str):
        """
        向知识库中添加一段文档文本，为其生成向量并保存
        """
        text_clean = text.strip()
        if not text_clean:
            return
            
        # 重复性检查，避免重复将完全相同的文本写入向量数据库
        if any(item["text"] == text_clean for item in self.knowledge_base):
            return
            
        print(f"📝 正在为文本片段生成向量并加入知识库: {text_clean[:30]}...")
        vector = self.get_embedding(text_clean)
        if vector:
            self.knowledge_base.append({
                "text": text_clean,
                "vector": vector
            })

    def _split_markdown(self, text: str, chunk_size: int = 600, overlap: int = 100) -> list:
        """
        简单的文本切片方法：将超长文档切分成 600 字左右的小片段，方便向量检索。
        保留 100 字的重叠区（overlap）能防止长句在切分处断开丢失语义。
        """
        chunks = []
        paragraphs = text.split("\n\n")
        current_chunk = ""
        
        for p in paragraphs:
            p = p.strip()
            if not p:
                continue
            if len(current_chunk) + len(p) < chunk_size:
                current_chunk += p + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = p + "\n\n"
                
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        # 滑动窗口硬切分：万一有单个段落字数极其庞大，用固定字数切割它
        final_chunks = []
        for c in chunks:
            if len(c) > chunk_size + 200:
                start = 0
                while start < len(c):
                    final_chunks.append(c[start:start + chunk_size])
                    start += chunk_size - overlap
            else:
                final_chunks.append(c)
                
        return final_chunks

    def ingest_file_via_mineru(self, file_path: str):
        """
        核心方法：利用 MinerU 提取工具导入文档（高保真转换）。
        加入了对“已导入文件”的过滤——若文件已经学习过，直接跳过生成向量，绝不消耗任何 Token。
        """
        if not os.path.exists(file_path):
            print(f"❌ 找不到需要导入的参考文件: {file_path}")
            return
            
        file_name = os.path.basename(file_path)
        # 📌 关键性能优化：若文件曾被加载过，直接跳过提取和 Embedding 生成！
        if file_name in self.ingested_files:
            print(f"⚡ 跳过: 风格样本 '{file_name}' 之前已经学习并存入本地数据库，直接重用。")
            return
            
        temp_dir = "./data/output/mineru_temp"
        os.makedirs(temp_dir, exist_ok=True)
        
        # MinerU-open-api 命令行调用参数：将文件转换并将 Markdown 输出到 temp_dir
        cmd = ["mineru-open-api", "flash-extract", file_path, "-o", temp_dir]
        print(f"📄 尝试调用 MinerU 命令行提取新样本: {' '.join(cmd)}")
        
        success = False
        content = ""
        
        try:
            # 开启 shell=True 兼容 Windows PATH 下的 exe/cmd 命令
            result = subprocess.run(
                cmd, 
                capture_output=True, 
                text=True, 
                shell=True, 
                encoding="utf-8", 
                errors="ignore"
            )
            if result.returncode == 0:
                # 定位 MinerU 转换生成的 Markdown 文件（通常后缀为 .md）
                md_files = glob.glob(os.path.join(temp_dir, "**", "*.md"), recursive=True)
                if not md_files:
                    md_files = glob.glob(os.path.join(temp_dir, "*.md"))
                    
                if md_files:
                    md_path = md_files[0]
                    print(f"💡 成功定位 MinerU 转换的 Markdown 文件: {md_path}")
                    with open(md_path, "r", encoding="utf-8") as f:
                        content = f.read()
                    success = True
                else:
                    print("⚠️ MinerU 转换成功但没找到输出的 .md 文件，准备降级为普通读取...")
            else:
                print(f"⚠️ MinerU 提取返回失败 (退出码: {result.returncode})，准备降级为普通读取...")
                
        except Exception as e:
            print(f"💡 提示：未在您的系统里检测到 MinerU 命令行工具或执行出错。系统已自动切换为内置的 Python 读取器进行解析。（错误信息: {e}）")
            
        # 降级兜底方案：如果 MinerU 不可用或提取失败，直接使用 Python 库提取
        if not success:
            ext = os.path.splitext(file_path)[1].lower()
            print(f"🔄 启动原生 Python 读取器提取 {ext} 文件内容...")
            try:
                if ext == ".docx":
                    doc = docx.Document(file_path)
                    text_parts = []
                    # 读段落
                    for p in doc.paragraphs:
                        if p.text.strip():
                            text_parts.append(p.text.strip())
                    # 读表格
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text and (not text_parts or text_parts[-1] != cell_text):
                                    text_parts.append(cell_text)
                    content = "\n\n".join(text_parts)
                else:
                    # 默认按照纯文本读取（适用于 .txt 和 .md）
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                success = True
            except Exception as ex:
                print(f"❌ 兜底方案文件解析失败: {ex}")
                
        # 3. 对提取出的文字做语义分块并生成嵌入向量
        if success and content.strip():
            chunks = self._split_markdown(content, chunk_size=600)
            print(f"✂️ 文档已被划分为 {len(chunks)} 个片段，正在生成嵌入向量并写入向量库...")
            for chunk in chunks:
                self.add_document(chunk)
            
            # 将该文件名记录到已处理列表中
            self.ingested_files.append(file_name)
            self.save_db()
            print("🎉 新学习文档已成功导入 RAG 向量知识库！")
            
        # 4. 清理临时生成的 Markdown 文件和目录
        shutil.rmtree(temp_dir, ignore_errors=True)

    def _cosine_similarity(self, v1: list, v2: list) -> float:
        """
        纯 Python 计算两个向量的余弦相似度（无需安装 NumPy）
        公式：Similarity = (A · B) / (||A|| * ||B||)
        """
        if not v1 or not v2 or len(v1) != len(v2):
            return 0.0
        # 1. 向量点积（A · B）
        dot_product = sum(a * b for a, b in zip(v1, v2))
        # 2. 向量 A 的 L2 范数（模长）
        norm_a = sum(a * a for a in v1) ** 0.5
        # 3. 向量 B 的 L2 范数（模长）
        norm_b = sum(b * b for b in v2) ** 0.5
        
        # 避免分母为 0
        if norm_a * norm_b > 0:
            return dot_product / (norm_a * norm_b)
        return 0.0

    def retrieve(self, query: str, top_k: int = 2) -> str:
        """
        根据检索问题，计算与知识库中所有片段的余弦相似度，并返回最相关的 top_k 段背景上下文
        """
        if not self.knowledge_base:
            return ""
            
        # 获取检索问题的向量
        query_vector = self.get_embedding(query)
        if not query_vector:
            return ""
            
        scored_chunks = []
        for chunk in self.knowledge_base:
            # 计算纯 Python 余弦相似度
            similarity = self._cosine_similarity(query_vector, chunk["vector"])
            scored_chunks.append((similarity, chunk["text"]))
            
        # 按照相似度从高到低排序
        scored_chunks.sort(key=lambda x: x[0], reverse=True)
        
        retrieved_texts = []
        for i in range(min(top_k, len(scored_chunks))):
            score, text = scored_chunks[i]
            # 过滤掉低相关性的内容（阈值设为 0.35）
            if score > 0.35: 
                retrieved_texts.append(f"[参考资料片段 {i+1}](相似度: {score:.2f}):\n{text}")
                
        return "\n\n".join(retrieved_texts)

# 本地单体测试
if __name__ == "__main__":
    try:
        rag = RAGEngine()
        sample_docs = [
            "大一学生写报告风格：通常句子比较通俗，喜欢用'我觉得'，段落结构清晰，不会堆砌太多高深的学术词汇。",
            "Matplotlib 中文乱码通用解决方案：在绘图前加入代码`plt.rcParams['font.sans-serif'] = ['SimHei']` 设置默认字体为黑体，并设置`plt.rcParams['axes.unicode_minus'] = False` 解决负号显示问题。"
        ]
        if not rag.knowledge_base:
            print("🌱 正在初始化测试知识库并导入参考资料...")
            for doc in sample_docs:
                rag.add_document(doc)
            rag.save_db()
            
        query = "如何解决 matplotlib 绘图时的中文乱码？"
        print(f"\n🔍 模拟检索提问: '{query}'")
        context = rag.retrieve(query, top_k=2)
        print("\n--- 检索到的上下文 ---")
        print(context)
    except Exception as e:
        print(f"RAG 初始化/运行测试出错: {e}")