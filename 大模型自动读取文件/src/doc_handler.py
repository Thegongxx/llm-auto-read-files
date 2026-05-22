# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

class DocumentHandler:
    def __init__(self, temp_dir="./data/output"):
        """
        初始化 Word 文档处理器
        """
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)

    def read_docx_text(self, file_path: str) -> str:
        """
        读取一个 .docx 文件的所有文本（包括段落和表格里的文字）
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"未找到 Word 文件: {file_path}")
            
        doc = docx.Document(file_path)
        full_text = []
        
        # 1. 提取所有段落的文字
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
                
        # 2. 提取所有表格中的文字
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # 避免把完全重复的单元格文字重复添加
                    if cell_text and (not full_text or full_text[-1] != cell_text):
                        full_text.append(cell_text)
                        
        return "\n".join(full_text)

    def format_paragraph_to_songti5(self, paragraph, font_name="宋体", size_pt=10.5):
        """
        将段落文字格式化为“五号宋体字”，并调整段后距和行距。
        
        💡 大白话科普：
        为什么这下面有几行看起来很高深难懂的“_r.get_or_add_rPr()”？
        因为 python-docx 库本身有个陈年 Bug —— 它在设置字体时，只能影响西文字符（英文和数字），
        遇到中文时，Word 会自动使用系统默认字体，导致你在代码里设了“宋体”，导出的 Word 里却还是别扭的微软雅黑。
        为了强行修复这个 Bug，我们通过底层的 XML 协议节点，把 ASCII (西文)、hAnsi (符号) 和 eastAsia (东亚中文)
        三处的字体全部强行捆绑为“宋体”，这样就能做到百分之百完美显示中文宋体啦！
        """
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.15
        
        # 遍历段落中每一小块文字运行(run)，并应用样式
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            
            # 以下三行是强行修复中文无法显示宋体的底层 XML 操作
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)

    def clear_cell_content(self, cell):
        """
        清空指定 Word 表格单元格里的原有内容，以便重新写入
        """
        p_list = list(cell.paragraphs)
        for i, p in enumerate(p_list):
            if i == 0:
                p.text = ""  # 第一个段落保留，仅清空文字
            else:
                # 剩下的段落直接移除，避免空白行堆叠导致占页码
                p._element.getparent().remove(p._element)

    def _clear_table_borders(self, table):
        """
        通过底层 XML 将嵌套表格的所有边框隐藏，使其变成“隐形网格”。
        这样在排版多张图时，图片能非常整齐地并排陈列，且没有难看的表格线。
        """
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        else:
            tblBorders.clear()
            
        # 设置六个方向的边框为 none
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)

    def write_code_to_cell(self, cell, code_text: str):
        """
        将生成的 Python 实验代码写入 Word 单元格中。
        这里行间距设置为紧凑的 1.0 倍，能有效地将长代码控制在指定的 3 页纸内。
        """
        self.clear_cell_content(cell)
        p = cell.paragraphs[0]
        
        lines = code_text.splitlines()
        for idx, line in enumerate(lines):
            # 将代码按行写入段落
            if idx == 0:
                p.text = line
            else:
                p = cell.add_paragraph(line)
            
            # 设置紧密排版格式
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            if not p.runs:
                p.add_run()
            self.format_paragraph_to_songti5(p, font_name="宋体", size_pt=10.5)

    def write_results_to_cell(self, cell, text_analysis: str, image_paths: list):
        """
        将大模型生成的文字分析和实验结果图表填入 Word 中。
        - 文字分析以 1.15 倍行距排版。
        - 多张图片使用“隐形网格”自适应宽度排列，防止图片过大导致打印页数超出 2 页的要求。
        """
        self.clear_cell_content(cell)
        
        # 1. 写入实验结果文字分析
        lines = text_analysis.splitlines()
        first_p = True
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            if first_p:
                p_analysis = cell.paragraphs[0]
                p_analysis.text = line_str
                first_p = False
            else:
                p_analysis = cell.add_paragraph(line_str)
            self.format_paragraph_to_songti5(p_analysis)
        
        # 2. 如果存在图片，自动在文字下方排列图片
        if image_paths:
            # 过滤出确实存在于本地的图片路径
            valid_paths = [p for p in image_paths if os.path.exists(p)]
            if not valid_paths:
                return
                
            num_images = len(valid_paths)
            if num_images == 1:
                # 只有一张图时，无需网格直接居中插入，限制宽度为 4.5 英寸
                p_img = cell.add_paragraph()
                p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after = Pt(6)
                run = p_img.add_run()
                run.add_picture(valid_paths[0], width=Inches(4.5))
                
                # 图注（小五号 9 pt）
                p_caption = cell.add_paragraph("图1 实验结果图")
                p_caption.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                self.format_paragraph_to_songti5(p_caption, size_pt=9.0)
            else:
                # 多张图时，创建 rows x 2 的隐形表格，分两列网格展示，美观且节省空间
                rows = (num_images + 1) // 2
                nested_table = cell.add_table(rows=rows, cols=2)
                nested_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 彻底清除网格线，使其在 Word 里表现为隐形
                self._clear_table_borders(nested_table)
                
                # 初始化嵌套单元格并清除冗余空行
                for r in range(rows):
                    for c in range(2):
                        grid_cell = nested_table.cell(r, c)
                        self.clear_cell_content(grid_cell)
                        grid_cell.width = Inches(2.8)
                
                # 将图片和对应的图注填入隐形单元格中
                for idx, img_path in enumerate(valid_paths):
                    r_idx = idx // 2
                    c_idx = idx % 2
                    grid_cell = nested_table.cell(r_idx, c_idx)
                    
                    grid_p = grid_cell.paragraphs[0]
                    grid_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    
                    run = grid_p.add_run()
                    # 限制单张图宽为 2.8 英寸，两张并排刚好贴合 A4 纸宽度
                    run.add_picture(img_path, width=Inches(2.8))
                    
                    # 填入图注
                    fig_num = idx + 1
                    p_caption = grid_cell.add_paragraph(f"图{fig_num} 实验结果图")
                    p_caption.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    p_caption.paragraph_format.space_before = Pt(2)
                    p_caption.paragraph_format.space_after = Pt(2)
                    self.format_paragraph_to_songti5(p_caption, size_pt=9.0)

    def apply_styles_to_document(self, doc_path: str, font_name="宋体", size_pt=10.5):
        """
        保修功能：全局扫描并强制统一整个 Word 文档中所有文本的字体为五号宋体
        """
        doc = docx.Document(doc_path)
        
        def format_run(run):
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)

        # 1. 扫描正文段落
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                format_run(run)
                
        # 2. 扫描表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            format_run(run)
                            
        doc.save(doc_path)
        print(f"📌 全局字体强制统一完成（已统一设为五号宋体）。")
